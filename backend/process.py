# Importing Libraries and Google AI API Key
# import google
import google.generativeai as genai
# from google.generativeai import caching
from dotenv import load_dotenv
# import pathlib
import os
import time
# import datetime
import re
import json
import cv2

# !pip install docx
# !pip install --upgrade python-docx

# import pandas as pd
# import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

load_dotenv()

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
genai.configure(api_key=GEMINI_API_KEY)

generation_config = {
    "temperature": 1,
    "top_p": 0.95,
    "top_k": 40,
    "max_output_tokens": 8192,
    "response_mime_type": "text/plain",
}

# System Instuction for generating the process documentation
generate_docs_systemPrompt = """You are a Business Analyst tasked with reviewing a process recording from the Subject Matter Expert (SME) in the form of a video. Your objective is to carefully analyze the video and extract a detailed, step-by-step outline of the process presented. The video may not cover the process end-to-end, so you need to assess both the explicit steps presented and any references the SME makes to previous steps. Your outline should be clear, precise, and suitable for inclusion in formal documentation, such as a Process Definition Document (PDD). Ensure that each step is detailed, any business exceptions are noted, and the process is presented in the order it is executed. Pay attention to the narrator’s comments to identify any transitions or additional information. The structure of the output documentation should include the following sections: 1. Process Name Provide the name of the process being described. 2. Short Process Description Offer a brief summary of the process. 3. List of Applications Utilized This should be a table that includes the following details for each application used in the process: - The name of the application - The type of the application (e.g., web application, desktop application) - The URL of the application, if applicable Ensure both web and desktop applications are identified. 4. List of Steps - Provide a detailed, step-by-step description of the process in the order the steps are executed. - Steps should be listed as they were presented in the video. - Each interaction with the user interface (UI) must be documented. - Document each described or presented data transformation. - Use the following numbering format: - Example: 1.0 Group of steps 1.1 First step in the group 1.2 Second step in the group - Steps should specify the UI element the user interacts with or the calculation logic described. - First step the group should specify the application name that the user interact with. 5. Exception Handling Describe any exceptions in the process and how they should be handled. 6. Requires Clarification List any questions you have for the SME or aspects of the process that require further clarification. Provide the output in the following JSON format: { "process_name": "[The name of the process based on the video content]", "short_process_description": "[The short process description based on the video content]", "list_of_applications": [ { "application_name": "[Name of the application]", "type": "[Type of the application, e.g., web/desktop]", "url": "[URL of the application, if applicable]" }, { "application_name": "[Name of the application]", "type": "[Type of the application, e.g., web/desktop]", "url": "[URL of the application, if applicable]" } ], "list_of_steps": [ { "group_name": "[Description of the group of steps]", "numbering": "1.0", "time_stamp": "[Timestamp from the video when this step is executed]", "sub_steps": [ { "step": "[Description of the sub-step]", "numbering": "1.1", "time_stamp": "[Timestamp from the video when this step is executed]" }, { "step": "[Description of the sub-step]", "numbering": "1.2", "time_stamp": "[Timestamp from the video when this step is executed]" } ] } ], "exceptions": [ { "exception": "[Exception name]", "description": "[Exception description]" }, { "exception": "[Exception name]", "description": "[Exception description]" }, { "exception": "[Exception name]", "description": "[Exception description]" } ], "clarifications": [ "[Required clarification or question]", "[Required clarification or question]", "[Required clarification or question]" ] } """

model = genai.GenerativeModel(
    model_name="gemini-1.5-flash",
    generation_config=generation_config,
    system_instruction=generate_docs_systemPrompt)


# Functions for managing the video files upload
def upload_to_gemini(path, mime_type=None):
    file = genai.upload_file(path, mime_type=mime_type)
    print(f"Uploaded file '{file.display_name}' as: {file.uri}")
    return file

def wait_for_files_active(files):
    print("⏳ Waiting for file processing...")
    for name in (file.name for file in files):
        file = genai.get_file(name)
        while file.state.name == "PROCESSING":
            # print(".", end="", flush=True)
            time.sleep(10)
            file = genai.get_file(name)
        if file.state.name != "ACTIVE":
            raise Exception(f"File {file.name} failed to process")
    print("✅ All files ready")

# Video frame extraction function
def extract_frame(video_path, timestamp, output_path):
    try:
        # Parse timestamp to seconds
        parts = list(map(int, timestamp.split(':')))
        if len(parts) == 3:
            h, m, s = parts
            total_seconds = h * 3600 + m * 60 + s
        elif len(parts) == 2:
            m, s = parts
            total_seconds = m * 60 + s
        else:
            total_seconds = parts[0]

        cap = cv2.VideoCapture(video_path)
        cap.set(cv2.CAP_PROP_POS_MSEC, total_seconds * 1000)
        success, frame = cap.read()
        if success:
            cv2.imwrite(output_path, frame)
        return success
    except Exception as e:
        print(f"Error extracting frame: {e}")
        return False
    finally:
        if 'cap' in locals():
            cap.release()

# Document creation functions
def set_table_borders(table, border_color="auto"):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')

            for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border_elem = OxmlElement(f'w:{border}')
                border_elem.set(qn('w:val'), 'single')
                border_elem.set(qn('w:sz'), '4')
                border_elem.set(qn('w:space'), '0')
                if border_color != "auto":
                    border_elem.set(qn('w:color'), border_color)
                tcBorders.append(border_elem)
            tcPr.append(tcBorders)

def create_document(json_data, video_path):
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'

    # Document header (unchanged)
    doc.add_heading('Process Name: ' + json_data["process_name"], level=1)
    doc.add_paragraph(json_data["short_process_description"])

    # Applications table (unchanged)
    doc.add_heading('List of applications', level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text = 'Application Name', 'Type', 'URL'
    for cell in hdr_cells: cell.paragraphs[0].runs[0].bold = True

    for app in json_data["list_of_applications"]:
        row_cells = table.add_row().cells
        row_cells[0].text = app['application_name']
        row_cells[1].text = app['type']
        row_cells[2].text = app.get('url', '')

    # Modified steps table with middle cell screenshots
    doc.add_heading('List of steps', level=2)
    for step_group in json_data["list_of_steps"]:
        # Step group header (unchanged)
        table = doc.add_table(rows=1, cols=3)
        for i, width in enumerate([Inches(0.5), Inches(4.5), Inches(1.5)]):
            table.columns[i].width = width

        hdr_cell = table.rows[0].cells[0].merge(table.rows[0].cells[2])
        hdr_run = hdr_cell.paragraphs[0].add_run(f"{step_group['numbering']} {step_group['group_name']}")
        hdr_run.bold, hdr_run.font.color.rgb = True, RGBColor(54, 95, 145)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D1E2F8')
        hdr_cell._element.tcPr.append(shading_elm)

        # Modified steps and screenshots
        for sub_step in step_group["sub_steps"]:
            # Step row (unchanged)
            row_cells = table.add_row().cells
            row_cells[0].text = sub_step['numbering']
            row_cells[1].text = sub_step['step']
            row_cells[2].text = sub_step['time_stamp']

            # Screenshot row - modified to use middle cell only
            timestamp = sub_step['time_stamp']
            screenshot_path = f"screenshot_{sub_step['numbering'].replace('.', '_')}.png"

            if extract_frame(video_path, timestamp, screenshot_path):
                # Add new row with three separate cells
                img_row = table.add_row()
                # Leave first and third cells empty
                img_row.cells[0].text = ''
                img_row.cells[2].text = ''

                # Add image to middle cell
                middle_cell = img_row.cells[1]
                paragraph = middle_cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(screenshot_path, width=Inches(4))  # Slightly smaller than cell width
                os.remove(screenshot_path)
            else:
                # Error message in middle cell
                error_row = table.add_row()
                error_row.cells[0].text = ''
                error_row.cells[2].text = ''
                error_row.cells[1].text = "Screenshot unavailable"
                error_row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        set_table_borders(table)

    # Exceptions and clarifications
    doc.add_heading('Exceptions', level=2)
    for exception in json_data["exceptions"]:
        doc.add_paragraph(f"Exception: {exception['exception']}", style='List Bullet')
        doc.add_paragraph(f"Description: {exception['description']}")

    doc.add_heading('Requires Clarification', level=2)
    for clarification in json_data["clarifications"]:
        doc.add_paragraph(clarification, style='List Bullet')

    # # Save document
    # doc_name = f"{json_data['process_name'].replace(' ', '_')}.docx"
    # doc.save(doc_name)
    # return doc_name

    # Ensure 'processed' folder exists
    output_folder = "output"
    os.makedirs(output_folder, exist_ok=True)

    # Save document
    doc_name = f"{json_data['process_name'].replace(' ', '_')}.docx"
    doc_path = os.path.join(output_folder, doc_name)  # Save inside 'processed' folder
    doc.save(doc_path)

    return doc_path


def videoToDoc(video_path, language="English"):
    # Upload and process video
    # video_path = "./E2E Process - Offer Acceptation.mp4"
    files = [upload_to_gemini(video_path, "video/mp4")]
    wait_for_files_active(files)

    # Append language instruction to the existing system prompt
    updated_instruction = generate_docs_systemPrompt + f"\n\nPlease provide response in {language}."
    
    # Create a new model instance with the updated system instruction
    model_with_language = genai.GenerativeModel(
        model_name="gemini-1.5-flash",
        generation_config=generation_config,
        system_instruction=updated_instruction
    )

    # Generate documentation
    chat_session = model_with_language.start_chat(history=[{"role": "user", "parts": files}])
    response = chat_session.send_message("Process the attached video.")

    # Parse response
    json_match = re.search(r"\{.*\}", response.text, re.DOTALL)
    if not json_match:
        raise ValueError("No valid JSON found in response")

    json_data = json.loads(json_match.group(0))

    # Create document with screenshots
    doc_path = create_document(json_data, video_path)
    print(f"Document generated successfully: {doc_path}")
    return doc_path